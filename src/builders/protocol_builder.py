"""
Assembles the final protocol DOCX from extracted fragments.

Strategy: create a new document, copy styles from the first source DOCX,
add header paragraphs, then insert each fragment's raw XML nodes in order.
"""
import copy
import zipfile
from typing import List, Optional
from lxml import etree

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from ..models.fragment import Fragment, FragmentStatus
from ..models.protocol_header import ProtocolHeader
from ..utils.anchors import point_sort_key, W_NS

_W = f'{{{W_NS}}}'


def build_protocol(
    fragments: List[Fragment],
    header: ProtocolHeader,
    output_path: str,
) -> str:
    doc = Document()

    # Copy styles from the first DOCX source so referenced styles resolve
    first_docx = next(
        (f.file_path for f in fragments if f.source_type == 'docx' and f.xml_nodes),
        None,
    )
    if first_docx:
        _inject_styles(doc, first_docx)

    _setup_page(doc)
    _add_header(doc, header)

    _bad = (FragmentStatus.NOT_FOUND, FragmentStatus.ERROR)
    included = [f for f in fragments if f.include_in_protocol and f.status not in _bad]
    problematic = [f for f in fragments if f.include_in_protocol and f.status in _bad]
    included = _sort_by_point(included)

    body = doc.element.body
    sect_pr = body.find(qn('w:sectPr'))

    for frag in included:
        frag.assigned_number = frag.original_number
        if frag.source_type == 'docx' and frag.xml_nodes:
            nodes = [copy.deepcopy(n) for n in frag.xml_nodes]
            _scale_fonts(nodes, delta=4)
            _insert_nodes(body, nodes, sect_pr)
        else:
            _add_text_fragment(doc, frag, sect_pr)

        # Blank line between fragments
        _insert_nodes(body, [_blank_para()], sect_pr)

    if problematic:
        _add_problematic_section(doc, problematic, sect_pr)

    doc.save(output_path)
    return output_path


# ─── helpers ──────────────────────────────────────────────────────────────────

def _sort_by_point(fragments: List[Fragment]) -> List[Fragment]:
    """Sort by original point number (1.6 before 1.30); keep user order for ties."""
    return [
        frag for _, frag in sorted(
            enumerate(fragments),
            key=lambda pair: (point_sort_key(pair[1].original_number), pair[0]),
        )
    ]


def _insert_nodes(body, nodes, sect_pr):
    if sect_pr is not None:
        idx = list(body).index(sect_pr)
        for node in nodes:
            body.insert(idx, node)
            idx += 1
    else:
        for node in nodes:
            body.append(node)


def _blank_para() -> etree._Element:
    p = etree.Element(f'{_W}p')
    pPr = etree.SubElement(p, f'{_W}pPr')
    spacing = etree.SubElement(pPr, f'{_W}spacing')
    spacing.set(f'{_W}before', '120')
    spacing.set(f'{_W}after', '120')
    return p


def _scale_fonts(nodes: list, delta: int = 4):
    """Increase <w:sz> and <w:szCs> values by delta half-points in all nodes."""
    for node in nodes:
        for tag in (f'{_W}sz', f'{_W}szCs'):
            for elem in node.iter(tag):
                val = elem.get(f'{_W}val')
                if val and val.isdigit():
                    elem.set(f'{_W}val', str(int(val) + delta))


def _setup_page(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)


def _add_header(doc: Document, hdr: ProtocolHeader):
    def add(text: str, bold: bool = False, center: bool = False, size: int = 12):
        if not text:
            return
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.bold = bold

    if hdr.organization:
        add(hdr.organization, bold=True, center=True, size=14)
        doc.add_paragraph()

    title_parts = ['ПРОТОКОЛ']
    if hdr.protocol_number:
        title_parts.append(f'№ {hdr.protocol_number}')
    if hdr.protocol_date:
        title_parts.append(f'от {hdr.protocol_date}')
    add(' '.join(title_parts), bold=True, center=True, size=14)

    if hdr.session_name:
        add(hdr.session_name, center=True)

    doc.add_paragraph()

    if hdr.participants:
        add('Присутствовали:', bold=True)
        for line in hdr.participants.splitlines():
            add(line.strip())
        doc.add_paragraph()

    if hdr.secretary:
        add(f'Секретарь: {hdr.secretary}')
        doc.add_paragraph()

    if hdr.quorum_text:
        add(hdr.quorum_text)
        doc.add_paragraph()

    if hdr.agenda_item:
        add('ПОВЕСТКА ДНЯ:', bold=True)
        add(hdr.agenda_item)
        doc.add_paragraph()


def _add_text_fragment(doc: Document, frag: Fragment, sect_pr):
    body = doc.element.body

    lines = frag.text_preview.splitlines()
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)


def _add_problematic_section(doc: Document, fragments: List[Fragment], sect_pr):
    body = doc.element.body

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run('Файлы, требующие ручной проверки:')
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    for frag in fragments:
        import os
        name = os.path.basename(frag.file_path)
        status_label = frag.status.value
        msg = frag.error_message or '; '.join(frag.warnings)
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        r = p2.add_run(f'• {name} — {status_label}')
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        if msg:
            r2 = p2.add_run(f' ({msg})')
            r2.font.name = 'Times New Roman'
            r2.font.size = Pt(10)


def _inject_styles(doc: Document, source_docx: str):
    """Replace the target document's styles.xml with the source's."""
    try:
        with zipfile.ZipFile(source_docx, 'r') as z:
            styles_xml = z.read('word/styles.xml')
        new_styles = etree.fromstring(styles_xml)
        doc.element.body.getparent().replace(
            doc.element.body.getparent().find(qn('w:styles')),
            new_styles,
        )
    except Exception:
        pass
